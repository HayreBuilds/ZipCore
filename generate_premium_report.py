"""
Premium Report Generator — Huffman Coding Assignment II
Generates a professional, polished DOCX report.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color Palette ──
NAVY      = RGBColor(10, 25, 49)
DARK_BLUE = RGBColor(20, 50, 90)
MID_BLUE  = RGBColor(40, 80, 130)
ACCENT    = RGBColor(0, 122, 204)
SLATE     = RGBColor(90, 110, 135)
BODY_CLR  = RGBColor(30, 30, 42)
WHITE     = RGBColor(255, 255, 255)
LIGHT_BG  = RGBColor(240, 244, 248)
CODE_BG   = RGBColor(30, 30, 46)
CODE_FG   = RGBColor(205, 214, 244)
GREEN_ACC = RGBColor(0, 180, 120)
RED_ACC   = RGBColor(220, 60, 60)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "huffman_output")
REPORT_PATH = os.path.join(os.path.dirname(__file__), "Huffman_Coding_Premium_Report.docx")


def setup_styles(doc):
    """Configure all custom document styles."""
    # Normal body
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)
    s.font.color.rgb = BODY_CLR
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.line_spacing = 1.3
    s.paragraph_format.space_after = Pt(8)

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.color.rgb = NAVY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(28)
    h1.paragraph_format.space_after = Pt(10)

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.color.rgb = DARK_BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # List Bullet
    lb = doc.styles['List Bullet']
    lb.font.name = 'Calibri'
    lb.font.size = Pt(11)
    lb.font.color.rgb = BODY_CLR
    lb.paragraph_format.space_after = Pt(4)
    lb.paragraph_format.left_indent = Inches(0.4)

    return doc


def add_cover_page(doc):
    """Add a premium cover/title page."""
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ASSIGNMENT II")
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    r.font.bold = True
    r.font.letter_spacing = Pt(3)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    r2 = p2.add_run("Design and Implementation of\nLossless Text Compression\nUsing Huffman Coding")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(28)
    r2.font.color.rgb = NAVY
    r2.font.bold = True

    # Thin accent line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(16)
    line.paragraph_format.space_after = Pt(16)
    rl = line.add_run("━" * 40)
    rl.font.color.rgb = ACCENT
    rl.font.size = Pt(10)

    # Course info
    info_lines = [
        ("Course:", "Design and Analysis of Algorithms (DAA)"),
        ("Date:", "May 2026"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + " ")
        rl.font.name = 'Calibri'
        rl.font.size = Pt(12)
        rl.font.color.rgb = SLATE
        rl.bold = True
        rv = p.add_run(value)
        rv.font.name = 'Calibri'
        rv.font.size = Pt(12)
        rv.font.color.rgb = BODY_CLR

    doc.add_page_break()


def add_section_heading(doc, number, title):
    """Add a numbered section heading."""
    doc.add_heading(f"{number}. {title}", level=1)


def add_sub_heading(doc, number, title):
    """Add a numbered sub-heading."""
    doc.add_heading(f"{number} {title}", level=2)


def add_body(doc, text):
    """Add body paragraph."""
    doc.add_paragraph(text)


def add_bullet(doc, text):
    """Add a bullet point."""
    doc.add_paragraph(text, style='List Bullet')


def shade_cell(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_cell_text(cell, text, bold=False, color=None, size=Pt(10), center=False):
    """Style text inside a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = size
    r.font.bold = bold
    if color:
        r.font.color.rgb = color


def add_premium_table(doc, headers, rows):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "0A1931")
        style_cell_text(cell, h, bold=True, color=WHITE, size=Pt(10), center=True)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            style_cell_text(cell, str(val), color=BODY_CLR, size=Pt(10), center=True)

    doc.add_paragraph()


def add_code_block(doc, code_text, title=None):
    """Add a styled code block with dark background."""
    if title:
        p_title = doc.add_paragraph()
        r = p_title.add_run(f"  {title}")
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = ACCENT
        r.font.bold = True
        r.font.italic = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)

    # Add background shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E2E" w:val="clear"/>')
    pPr.append(shading)

    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(166, 218, 149)


def add_chart_image(doc, filename, caption=None):
    """Insert a chart image if it exists."""
    path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(9)
            r.font.color.rgb = SLATE
            r.font.italic = True


def build_report():
    """Build the complete premium report."""
    doc = Document()
    setup_styles(doc)

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_cover_page(doc)

    # ═══════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════
    add_section_heading(doc, 1, "Introduction")
    add_body(doc,
        "This report presents the design and implementation of a lossless text "
        "compression system using Huffman Coding. The system is developed for a "
        "national digital platform (e.g., an e-government system) that manages "
        "large volumes of textual data — including citizen records, reports, "
        "service logs, and legal documents."
    )
    add_body(doc,
        "As the platform scales, storage costs increase and data transmission "
        "becomes slower, especially in bandwidth-constrained environments. "
        "The implemented solution reduces storage space, optimizes data "
        "transmission, and preserves the exact original content through "
        "lossless compression. Given the sensitivity of government data, "
        "no information loss is acceptable, making lossy techniques unsuitable."
    )

    # ═══════════════════════════════════════════
    # 2. PROBLEM STATEMENT
    # ═══════════════════════════════════════════
    add_section_heading(doc, 2, "Problem Statement")
    add_body(doc,
        "This assignment focuses on designing and implementing a lossless "
        "compression system using Huffman Coding that:"
    )
    add_bullet(doc, "Minimizes the number of bits required to represent text")
    add_bullet(doc, "Ensures accurate reconstruction of the original data")
    add_bullet(doc, "Adapts to varying character frequency distributions")
    add_body(doc,
        "The challenge is to efficiently encode and decode text, evaluate "
        "compression performance, and analyze the suitability of Huffman "
        "Coding in real-world systems."
    )

    # ═══════════════════════════════════════════
    # 3. METHODOLOGY
    # ═══════════════════════════════════════════
    add_section_heading(doc, 3, "Methodology")

    add_sub_heading(doc, "3.1", "Algorithm Overview")
    add_body(doc,
        "Huffman Coding is a greedy algorithm that assigns variable-length "
        "prefix codes to characters based on their frequency of occurrence. "
        "Characters appearing more frequently receive shorter codes, while "
        "less frequent characters receive longer codes. This ensures no code "
        "is a prefix of another (prefix-free property), enabling unambiguous "
        "decoding."
    )
    add_body(doc, "The algorithm proceeds through the following steps:")
    add_bullet(doc,
        "Frequency Analysis — Count the occurrence of each character in the input text"
    )
    add_bullet(doc,
        "Priority Queue Construction — Create a min-heap of leaf nodes, one per unique character"
    )
    add_bullet(doc,
        "Tree Building — Repeatedly extract the two lowest-frequency nodes, merge them under "
        "a new internal node, and reinsert until one root remains"
    )
    add_bullet(doc,
        "Code Generation — Traverse the tree from root to leaves, assigning '0' for left "
        "branches and '1' for right branches"
    )
    add_bullet(doc,
        "Encoding — Replace each character in the input with its Huffman code"
    )
    add_bullet(doc,
        "Decoding — Traverse the Huffman tree bit-by-bit to reconstruct the original text"
    )

    add_sub_heading(doc, "3.2", "Implementation Details")
    add_body(doc,
        "The system was implemented in Python using the following key components:"
    )
    add_bullet(doc,
        "HuffmanNode class — Represents nodes in the Huffman tree with character, "
        "frequency, and child pointers"
    )
    add_bullet(doc,
        "HuffmanCoding class — Main class implementing frequency analysis, tree "
        "construction, code generation, encoding, and decoding"
    )
    add_bullet(doc,
        "File I/O — Reads text files, compresses to binary format, stores code "
        "tables as JSON, and reconstructs original files"
    )
    add_bullet(doc,
        "The heapq module provides an efficient min-heap priority queue for tree construction"
    )
    add_bullet(doc,
        "Bit padding ensures the encoded bitstream aligns to byte boundaries for file storage"
    )

    # ── Code Snippets ──
    add_sub_heading(doc, "3.3", "Core Implementation — Key Code Snippets")

    add_code_block(doc, """class HuffmanNode:
    \"\"\"A node in the Huffman Tree.\"\"\"
    def __init__(self, char=None, freq=0, left=None, right=None):
        self.char  = char
        self.freq  = freq
        self.left  = left
        self.right = right

    def __lt__(self, other):
        return self.freq < other.freq

    def is_leaf(self):
        return self.left is None and self.right is None""",
        title="▸ HuffmanNode — Tree Node Structure"
    )

    add_code_block(doc, """def build_huffman_tree(self, frequency_table):
    \"\"\"Construct a Huffman Tree using a min-heap.\"\"\"
    priority_queue = []
    for char, freq in frequency_table.items():
        heapq.heappush(priority_queue, HuffmanNode(char, freq))

    while len(priority_queue) > 1:
        left  = heapq.heappop(priority_queue)
        right = heapq.heappop(priority_queue)
        merged = HuffmanNode(
            freq=left.freq + right.freq,
            left=left, right=right
        )
        heapq.heappush(priority_queue, merged)

    self.root = priority_queue[0]
    return self.root""",
        title="▸ Huffman Tree Construction — Greedy Min-Heap"
    )

    add_code_block(doc, """def generate_codes(self, node=None, current_code=""):
    \"\"\"Generate prefix codes via DFS traversal.\"\"\"
    if node is None:
        node = self.root
    if node.is_leaf():
        code = current_code if current_code else "0"
        self.codes[node.char] = code
        self.reverse_codes[code] = node.char
        return
    self.generate_codes(node.left,  current_code + "0")
    self.generate_codes(node.right, current_code + "1")""",
        title="▸ Prefix Code Generation — DFS Traversal"
    )

    add_code_block(doc, """def encode(self, text):
    \"\"\"Compress text using Huffman codes.\"\"\"
    return "".join(self.codes[char] for char in text)

def decode(self, encoded_text):
    \"\"\"Decompress encoded bitstream to original text.\"\"\"
    decoded, current_code = [], ""
    for bit in encoded_text:
        current_code += bit
        if current_code in self.reverse_codes:
            decoded.append(self.reverse_codes[current_code])
            current_code = ""
    return "".join(decoded)""",
        title="▸ Encoding & Decoding — Compression Pipeline"
    )

    # ═══════════════════════════════════════════
    # 4. TEST DATASETS
    # ═══════════════════════════════════════════
    add_section_heading(doc, 4, "Test Datasets")
    add_body(doc,
        "Four distinct text datasets were created to evaluate compression "
        "performance across different text characteristics:"
    )
    add_premium_table(doc,
        ["Dataset", "Size (bytes)", "Unique Chars", "Description"],
        [
            ["Repetitive Text", "2,043", "34",
             "Citizen records with repetitive government terminology"],
            ["Legal Document", "1,984", "55",
             "Formal legal text — structured legislative vocabulary"],
            ["Service Logs", "1,434", "66",
             "System logs with timestamps and status codes"],
            ["Random Mixed", "457", "90",
             "Uniformly distributed characters, symbols, numbers"],
        ]
    )

    # ═══════════════════════════════════════════
    # 5. RESULTS AND ANALYSIS
    # ═══════════════════════════════════════════
    add_section_heading(doc, 5, "Results and Analysis")

    add_sub_heading(doc, "5.1", "Compression Performance Summary")
    add_body(doc,
        "The following table summarizes the compression performance across "
        "all four datasets. All datasets passed lossless verification — "
        "decoded output was byte-for-byte identical to the original input."
    )
    add_premium_table(doc,
        ["Dataset", "Original (B)", "Compressed (B)",
         "Ratio", "Savings (%)", "Bits/Char", "Verified"],
        [
            ["Repetitive Text", "2,043", "1,090",
             "1.87:1", "46.65%", "4.26", "✓ YES"],
            ["Legal Document", "1,984", "1,108",
             "1.79:1", "44.15%", "4.46", "✓ YES"],
            ["Service Logs", "1,434", "937",
             "1.53:1", "34.66%", "5.22", "✓ YES"],
            ["Random Mixed", "457", "366",
             "1.25:1", "19.91%", "6.38", "✓ YES"],
        ]
    )

    # Charts
    add_sub_heading(doc, "5.2", "Visual Analysis")

    add_chart_image(doc, "compression_comparison.png",
                    "Figure 1 — Original vs Compressed File Size Comparison")
    add_chart_image(doc, "space_savings.png",
                    "Figure 2 — Space Savings Percentage by Dataset Type")
    add_chart_image(doc, "bits_per_char.png",
                    "Figure 3 — Average Bits per Character (Huffman vs Fixed 8-bit ASCII)")
    add_chart_image(doc, "timing_comparison.png",
                    "Figure 4 — Encoding vs Decoding Time (milliseconds)")

    add_sub_heading(doc, "5.3", "Analysis of Results")
    add_body(doc, "Key findings from the compression analysis:")
    add_bullet(doc,
        "Repetitive Text achieved the highest compression (46.65% space savings) "
        "because repetitive patterns lead to skewed frequency distributions, "
        "allowing very short codes for common characters. With only 34 unique "
        "characters, the average code length was just 4.26 bits vs 8 bits in ASCII."
    )
    add_bullet(doc,
        "Legal Document achieved strong compression (44.15% savings) due to "
        "structured vocabulary. Despite 55 unique characters, common letters "
        "like spaces, 'e', 'n', and 't' dominated the frequency distribution."
    )
    add_bullet(doc,
        "Service Logs showed moderate compression (34.66% savings). While structured, "
        "the presence of numbers, timestamps, and identifiers (66 unique chars) "
        "created a more uniform distribution than natural language."
    )
    add_bullet(doc,
        "Random Mixed Text showed the lowest compression (19.91% savings) "
        "because its 90 unique characters had a nearly uniform frequency "
        "distribution — the worst case for Huffman Coding."
    )
    add_body(doc,
        "These results demonstrate that Huffman Coding is most effective when "
        "the character frequency distribution is highly skewed (non-uniform), "
        "which is typical of natural language text used in government systems."
    )

    add_sub_heading(doc, "5.4", "Lossless Verification")
    add_body(doc,
        "All four datasets passed lossless verification — the decoded output "
        "was byte-for-byte identical to the original input. This confirms that "
        "the implementation correctly preserves all original data, meeting the "
        "strict requirement for government data where no information loss is acceptable."
    )

    # ═══════════════════════════════════════════
    # 6. TIME COMPLEXITY ANALYSIS
    # ═══════════════════════════════════════════
    add_section_heading(doc, 6, "Time Complexity Analysis")
    add_body(doc, "The time complexity of the Huffman Coding algorithm:")

    add_premium_table(doc,
        ["Operation", "Complexity", "Description"],
        [
            ["Frequency Analysis", "O(n)", "Single pass through n characters"],
            ["Priority Queue Build", "O(k log k)", "Insert k unique chars into min-heap"],
            ["Tree Construction", "O(k log k)", "(k−1) extract-min + insert operations"],
            ["Code Generation", "O(k)", "Single DFS traversal of the tree"],
            ["Encoding", "O(n)", "Replace each of n chars with its code"],
            ["Decoding", "O(m)", "Traverse m bits of encoded text"],
        ]
    )

    add_body(doc,
        "Overall Encoding Complexity: O(n + k log k) ≈ O(n) since typically k ≪ n"
    )
    add_body(doc,
        "Overall Decoding Complexity: O(m) where m is the total number of encoded bits"
    )

    add_sub_heading(doc, "6.1", "Space Complexity")
    add_bullet(doc, "Huffman Tree: O(k) nodes where k = number of unique characters")
    add_bullet(doc, "Code Table: O(k × L) where L = average code length")
    add_bullet(doc, "Encoded Output: O(n × L_avg) — typically smaller than n × 8 (ASCII)")

    # ═══════════════════════════════════════════
    # 7. DISCUSSION
    # ═══════════════════════════════════════════
    add_section_heading(doc, 7, "Discussion — Suitability for Real-World Systems")

    add_sub_heading(doc, "7.1", "Advantages")
    add_bullet(doc, "Guaranteed lossless compression — essential for sensitive government data")
    add_bullet(doc, "Simple implementation with well-understood theoretical foundations")
    add_bullet(doc, "Effective for natural language text (40–50% space savings typical)")
    add_bullet(doc, "Fast encoding and decoding with linear time complexity")
    add_bullet(doc, "No patent restrictions — freely implementable")

    add_sub_heading(doc, "7.2", "Limitations")
    add_bullet(doc,
        "Requires two passes over the data (frequency analysis + encoding), "
        "or pre-computed frequency tables"
    )
    add_bullet(doc, "The code table must be stored or transmitted alongside compressed data")
    add_bullet(doc, "Less effective for data with uniform character distributions")
    add_bullet(doc,
        "Does not exploit higher-order patterns (word repetitions, phrases) — "
        "only single character frequencies"
    )

    add_sub_heading(doc, "7.3", "Recommendations for Production Deployment")
    add_body(doc,
        "For the national digital platform scenario, Huffman Coding provides a "
        "solid baseline compression mechanism. For production deployment, it "
        "could be enhanced by:"
    )
    add_bullet(doc, "Using adaptive Huffman coding to eliminate the two-pass requirement")
    add_bullet(doc,
        "Combining with dictionary-based methods (LZ77) as in the DEFLATE algorithm "
        "used by gzip/zlib"
    )
    add_bullet(doc, "Applying block-based compression for large files to manage memory efficiently")

    # ═══════════════════════════════════════════
    # 8. CONCLUSION
    # ═══════════════════════════════════════════
    add_section_heading(doc, 8, "Conclusion")
    add_body(doc,
        "This assignment successfully designed and implemented a lossless text "
        "compression system using Huffman Coding. The implementation demonstrates:"
    )
    add_bullet(doc, "Correct frequency analysis and Huffman tree construction")
    add_bullet(doc, "Optimal prefix-free code generation")
    add_bullet(doc, "Efficient encoding and decoding with verified lossless reconstruction")
    add_bullet(doc, "Compression ratios ranging from 1.25:1 to 1.87:1 depending on text characteristics")
    add_body(doc,
        "The system achieved up to 46.65% space savings on repetitive government-style "
        "text while guaranteeing perfect data reconstruction. The results confirm that "
        "Huffman Coding is well-suited for compressing natural language text in systems "
        "where data integrity is paramount, such as e-government platforms managing "
        "citizen records and legal documents."
    )
    add_body(doc,
        "The Python implementation (huffman_coding.py) serves as a complete, "
        "functional reference that can be adapted for production use."
    )

    # ═══════════════════════════════════════════
    # 9. REFERENCES
    # ═══════════════════════════════════════════
    add_section_heading(doc, 9, "References")
    refs = [
        "[1] Huffman, D.A. (1952). 'A Method for the Construction of "
        "Minimum-Redundancy Codes.' Proceedings of the IRE, 40(9), 1098–1101.",
        "[2] Cormen, T.H., Leiserson, C.E., Rivest, R.L., & Stein, C. (2009). "
        "Introduction to Algorithms (3rd ed.). MIT Press. Chapter 16: Greedy Algorithms.",
        "[3] Sayood, K. (2017). Introduction to Data Compression (5th ed.). Morgan Kaufmann.",
        "[4] Solomon, D. (2007). Data Compression: The Complete Reference (4th ed.). Springer.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════
    # HEADER & FOOTER
    # ═══════════════════════════════════════════
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Lossless Text Compression Using Huffman Coding  ·  DAA Assignment II")
    hr.font.name = 'Calibri'
    hr.font.size = Pt(8)
    hr.font.color.rgb = SLATE

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Design and Analysis of Algorithms — May 2026")
    fr.font.name = 'Calibri'
    fr.font.size = Pt(8)
    fr.font.color.rgb = SLATE

    doc.save(REPORT_PATH)
    print(f"✅ Premium report saved to: {REPORT_PATH}")


if __name__ == '__main__':
    build_report()
